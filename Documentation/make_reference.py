"""
Generates reference.docx — the Word style template used by Pandoc.
Run once:  python make_reference.py
Then Pandoc picks up styles from it automatically.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

NAVY  = RGBColor(0x1A, 0x5F, 0x7A)   # #1A5F7A
DARK  = RGBColor(0x14, 0x3C, 0x4E)   # darker navy for H1
GRAY  = RGBColor(0x44, 0x4C, 0x5C)   # body subtitle
BLACK = RGBColor(0x00, 0x00, 0x00)

doc = Document()

# ── Page layout ──────────────────────────────────────────────────────────────
for sec in doc.sections:
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.5)
    sec.page_width    = Cm(21.0)   # A4
    sec.page_height   = Cm(29.7)

# ── Helper: set font on a style ───────────────────────────────────────────────
def set_font(style, name, size_pt, bold=False, italic=False, color=None, space_before=0, space_after=0):
    f = style.font
    f.name        = name
    f.size        = Pt(size_pt)
    f.bold        = bold
    f.italic      = italic
    if color:
        f.color.rgb = color
    pf = style.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)

# ── Normal (body text) ────────────────────────────────────────────────────────
normal = doc.styles['Normal']
set_font(normal, 'Times New Roman', 12, color=BLACK, space_after=6)
normal.paragraph_format.line_spacing = Pt(18)

# ── Heading 1  (chapters) ─────────────────────────────────────────────────────
h1 = doc.styles['Heading 1']
set_font(h1, 'Arial', 18, bold=True, color=DARK, space_before=24, space_after=12)
h1.paragraph_format.keep_with_next = True
# Add bottom border
pPr = h1.element.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'),   'single')
bottom.set(qn('w:sz'),    '4')
bottom.set(qn('w:space'), '4')
bottom.set(qn('w:color'), '1A5F7A')
pBdr.append(bottom)
pPr.append(pBdr)

# ── Heading 2  (sections) ─────────────────────────────────────────────────────
h2 = doc.styles['Heading 2']
set_font(h2, 'Arial', 14, bold=True, color=NAVY, space_before=18, space_after=8)
h2.paragraph_format.keep_with_next = True

# ── Heading 3  (subsections) ──────────────────────────────────────────────────
h3 = doc.styles['Heading 3']
set_font(h3, 'Arial', 12, bold=True, color=NAVY, space_before=14, space_after=6)

# ── Heading 4  (subsubsections) ───────────────────────────────────────────────
h4 = doc.styles['Heading 4']
set_font(h4, 'Arial', 11, bold=True, italic=True, color=GRAY, space_before=10, space_after=4)

# ── Block Text (used for quotations / abstract) ───────────────────────────────
if 'Block Text' not in [s.name for s in doc.styles]:
    bt = doc.styles.add_style('Block Text', WD_STYLE_TYPE.PARAGRAPH)
else:
    bt = doc.styles['Block Text']
set_font(bt, 'Times New Roman', 11, color=GRAY, space_before=6, space_after=6)
bt.paragraph_format.left_indent  = Cm(1.2)
bt.paragraph_format.right_indent = Cm(1.2)

# ── Caption ───────────────────────────────────────────────────────────────────
cap = doc.styles['Caption']
set_font(cap, 'Arial', 10, italic=True, color=GRAY, space_before=4, space_after=10)
cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── Body Text / First Paragraph ───────────────────────────────────────────────
if 'Body Text' in [s.name for s in doc.styles]:
    bt2 = doc.styles['Body Text']
    set_font(bt2, 'Times New Roman', 12, color=BLACK, space_after=6)

# ── Code / Verbatim ───────────────────────────────────────────────────────────
if 'Verbatim Char' in [s.name for s in doc.styles]:
    vc = doc.styles['Verbatim Char']
    vc.font.name = 'Courier New'
    vc.font.size = Pt(10)

if 'Verbatim' in [s.name for s in doc.styles]:
    v = doc.styles['Verbatim']
    set_font(v, 'Courier New', 10, color=RGBColor(0x20, 0x30, 0x40), space_before=6, space_after=6)
    v.paragraph_format.left_indent = Cm(0.8)

# ── Table styles ──────────────────────────────────────────────────────────────
if 'Table' in [s.name for s in doc.styles]:
    tbl = doc.styles['Table']
    set_font(tbl, 'Arial', 10)

if 'Table Grid' in [s.name for s in doc.styles]:
    tg = doc.styles['Table Grid']
    set_font(tg, 'Arial', 10)

# ── Placeholder paragraph so file is valid ───────────────────────────────────
p = doc.add_paragraph('MediLingo AI — Reference Style Document', style='Heading 1')
doc.add_paragraph(
    'This file defines the Word styles used when Pandoc converts the LaTeX '
    'documentation to .docx format. Do not delete it.',
    style='Normal'
)
doc.add_paragraph('Chapter Heading (H1)', style='Heading 1')
doc.add_paragraph('Section Heading (H2)', style='Heading 2')
doc.add_paragraph('Subsection Heading (H3)', style='Heading 3')
doc.add_paragraph(
    'Body text in Times New Roman 12pt. Lorem ipsum dolor sit amet, '
    'consectetur adipiscing elit. Sed do eiusmod tempor incididunt.',
    style='Normal'
)
doc.add_paragraph('Figure 1: Example caption text', style='Caption')

doc.save('reference.docx')
print("reference.docx created successfully.")
